# src/utils/file_parser.py
#
# ─────────────────────────────────────────────────────────────────────────────
# WHAT IS THIS FILE?
#
# A single helper that extracts plain text from uploaded documents so the
# AI agent can read and reason about their contents directly — without
# storing them in the RAG knowledge base first.
#
# Supported formats:
#   .pdf   → PyMuPDF (fitz) with OCR fallback for scanned pages
#   .docx  → python-docx (reads paragraphs + tables)
#   anything else → decode as UTF-8 text
#
# WHY EXTRACT TEXT INSTEAD OF STORING IN RAG?
#   The RAG pipeline (rag.py) is designed for a *persistent* knowledge base:
#   chunk → embed → store in Milvus → retrieve later.
#   That's great for recurring reference docs, but for "I just uploaded a
#   contract, what does clause 5 say?" the overhead is unnecessary.
#
#   Here we extract raw text and pass it directly into Claude's context window.
#   Suitable for files up to ~80k tokens (≈ 60–80 pages of dense text).
#   For very large documents the caller should use RAG ingestion instead.
#
# CHARACTER LIMIT:
#   MAX_CHARS caps the extracted text at 80,000 characters to avoid
#   overflowing Claude's context window. That's roughly 20,000 tokens,
#   leaving plenty of room for history, system prompt, and the answer.
# ─────────────────────────────────────────────────────────────────────────────

import io
# io.BytesIO = wraps raw bytes in a file-like object.
# PyMuPDF and python-docx can open a BytesIO without writing a temp file to disk.

MAX_CHARS = 80_000
# Hard cap to keep Claude's context window safe.
# 80,000 characters ≈ 20,000 tokens (English prose averages ~4 chars/token).
# Increase this if you upgrade to a Claude model with a larger context window.


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract plain text from an uploaded file.

    Parameters
    ----------
    filename : str
        The original file name (e.g. "contract.pdf", "report.docx").
        Used only to determine the file type via its extension.

    content : bytes
        The raw binary content of the uploaded file.
        FastAPI's UploadFile gives you this via `await file.read()`.

    Returns
    -------
    str
        Plain text extracted from the document, capped at MAX_CHARS.
        May be truncated with a notice if the document is very large.

    Raises
    ------
    ValueError
        If the file is empty or the format is completely unreadable.

    PYTHON CONCEPT — `bytes`:
        `content: bytes` is raw binary data — not a string.
        You can't read it as text directly; you need a parser or .decode().
        Parsers like PyMuPDF and python-docx know the binary format and
        convert it to readable text for you.
    """
    if not content:
        raise ValueError("Uploaded file is empty.")

    name = filename.lower()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if name.endswith(".pdf"):
        return _extract_pdf(content)

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if name.endswith(".docx"):
        return _extract_docx(content)

    # ── Everything else: try UTF-8 text ──────────────────────────────────────
    # Handles .txt, .md, .csv, .json, .yaml, .html, .xml, etc.
    # errors="ignore" silently drops bytes that are not valid UTF-8
    # instead of raising a UnicodeDecodeError.
    text = content.decode("utf-8", errors="ignore").strip()
    if not text:
        raise ValueError("Could not extract any text from the file.")
    return _cap(text)


# ── PDF extraction ────────────────────────────────────────────────────────────

def _extract_pdf(content: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).

    Strategy per page:
      1. find_tables() → format as "col1 | col2 | col3" rows
      2. get_text()    → remaining paragraph / heading text
      3. OCR fallback  → if get_text() yields < 50 chars AND images exist
                         (indicates a scanned/image-only page)

    This mirrors the logic in src/routers/rag.py so both paths produce
    consistent output. Any improvements here should be reflected there too.
    """
    import fitz  # pymupdf — installed via `pip install pymupdf`

    pages_text: list[str] = []
    doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")
    # fitz.open(stream=...) opens the PDF from memory (no temp file needed).
    # filetype="pdf" tells fitz what parser to use when opening a stream.

    for page in doc:
        # `page` is a fitz.Page object — one page of the PDF.
        page_parts: list[str] = []

        # Step A: Extract tables as readable rows.
        # find_tables() was added in PyMuPDF 1.23; guard with try/except
        # so older installations still work (just without table formatting).
        try:
            for tab in page.find_tables():
                for row in tab.extract():
                    # row = list of cell strings (None for empty cells)
                    row_text = " | ".join(
                        str(cell).strip() for cell in row if cell
                    )
                    if row_text:
                        page_parts.append(row_text)
        except AttributeError:
            pass  # PyMuPDF < 1.23 — skip table extraction

        # Step B: Extract paragraph / heading text.
        plain = page.get_text() or ""
        # get_text() uses MuPDF's font decoder — handles Chinese CID fonts
        # correctly, unlike pdfminer.six which returns \u0001 placeholders.

        # Step C: OCR fallback for scanned / image-heavy pages.
        # Conditions: very short text AND the page contains embedded images.
        images_on_page = page.get_images()
        needs_ocr = (
            not plain.strip()
            or (images_on_page and len(plain.strip()) < 50)
        )

        if needs_ocr:
            try:
                import pytesseract
                from PIL import Image

                # Render the page at 2× zoom (higher DPI → better OCR accuracy).
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                # lang="chi_sim+eng" = Simplified Chinese + English language packs.
                # Install packs: brew install tesseract tesseract-lang (macOS)

                # Only replace plain if OCR found more content.
                if len(ocr_text.strip()) > len(plain.strip()):
                    plain = ocr_text
            except Exception:
                pass  # tesseract not installed — keep plain as-is

        if plain.strip():
            page_parts.append(plain.strip())

        # Join table rows + paragraph text for this page.
        pages_text.append("\n".join(page_parts))

    doc.close()

    full_text = "\n\n".join(pages_text).strip()
    if not full_text:
        raise ValueError("Could not extract any text from the PDF.")
    return _cap(full_text)


# ── DOCX extraction ───────────────────────────────────────────────────────────

def _extract_docx(content: bytes) -> str:
    """
    Extract text from a .docx file using python-docx.

    Reads:
      - Paragraphs (body text, headings, list items)
      - Tables (each cell joined with " | ", rows separated by newlines)

    python-docx is a pure-Python library — no external binary dependencies.
    Install: pip install python-docx
    """
    # PYTHON CONCEPT — lazy import:
    #   We import python-docx here instead of at the top of the file so
    #   the module loads even if python-docx is not installed.
    #   The ImportError only surfaces when someone actually uploads a .docx file.
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = Document(io.BytesIO(content))
    # Document(BytesIO(...)) opens the .docx from memory — no temp file needed.
    # A .docx file is actually a zip archive containing XML files.
    # python-docx unpacks the XML and exposes Python objects for each element.

    parts: list[str] = []

    # ── Paragraphs ────────────────────────────────────────────────────────────
    # doc.paragraphs = list of all paragraph objects in reading order.
    # Each paragraph has a .text attribute (plain string, no XML tags).
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # ── Tables ────────────────────────────────────────────────────────────────
    # doc.tables = list of all table objects in the document.
    # Each table has rows; each row has cells; each cell has .text.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise ValueError("Could not extract any text from the DOCX file.")
    return _cap(full_text)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _cap(text: str) -> str:
    """
    Truncate text to MAX_CHARS and append a notice if truncated.

    PYTHON CONCEPT — string slicing:
        text[:MAX_CHARS] returns the first MAX_CHARS characters of text.
        If len(text) <= MAX_CHARS, the whole string is returned unchanged.
        This is like JavaScript's text.slice(0, MAX_CHARS).
    """
    if len(text) <= MAX_CHARS:
        return text
    return (
        text[:MAX_CHARS]
        + f"\n\n[Document truncated at {MAX_CHARS:,} characters. "
        "Upload a shorter file or use RAG ingestion for large documents.]"
    )
