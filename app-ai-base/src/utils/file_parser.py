import io

MAX_CHARS = 80_000
# Hard cap to keep Claude's context window safe.
# 80,000 characters ≈ 20,000 tokens (English prose averages ~4 chars/token).


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract plain text from an uploaded file.

    Supported formats:
      .pdf  — PyMuPDF (fitz) with OCR fallback for scanned pages (pytesseract).
      .docx — python-docx (paragraphs + tables).
      other — UTF-8 decode (handles .txt, .md, .csv, .json, .yaml, etc.).

    Returns plain text capped at MAX_CHARS.
    Raises ValueError for empty files or completely unreadable content.
    """
    if not content:
        raise ValueError("Uploaded file is empty.")

    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)

    text = content.decode("utf-8", errors="ignore").strip()
    if not text:
        raise ValueError("Could not extract any text from the file.")
    return _cap(text)


def _extract_pdf(content: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).

    Per-page strategy:
      1. find_tables() — format table rows as "col1 | col2 | col3".
      2. get_text()    — remaining paragraph/heading text (handles Chinese CID fonts).
      3. OCR fallback  — pytesseract when get_text() < 50 chars and images are present.
    """
    import fitz  # pymupdf

    pages_text: list[str] = []
    doc = fitz.open(stream=io.BytesIO(content), filetype="pdf")

    for page in doc:
        page_parts: list[str] = []

        try:
            for tab in page.find_tables():
                for row in tab.extract():
                    row_text = " | ".join(str(cell).strip() for cell in row if cell)
                    if row_text:
                        page_parts.append(row_text)
        except AttributeError:
            pass  # PyMuPDF < 1.23 — skip table extraction

        plain          = page.get_text() or ""
        images_on_page = page.get_images()
        needs_ocr      = (not plain.strip()) or (images_on_page and len(plain.strip()) < 50)

        if needs_ocr:
            try:
                import pytesseract
                from PIL import Image

                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                if len(ocr_text.strip()) > len(plain.strip()):
                    plain = ocr_text
            except Exception:
                pass

        if plain.strip():
            page_parts.append(plain.strip())
        pages_text.append("\n".join(page_parts))

    doc.close()

    full_text = "\n\n".join(pages_text).strip()
    if not full_text:
        raise ValueError("Could not extract any text from the PDF.")
    return _cap(full_text)


def _extract_docx(content: bytes) -> str:
    """
    Extract text from a .docx file using python-docx.
    Reads paragraphs and tables; each table row joined with " | ".
    """
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx is not installed. Run: uv add python-docx")

    doc   = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise ValueError("Could not extract any text from the DOCX file.")
    return _cap(full_text)


def _cap(text: str) -> str:
    """Truncate text to MAX_CHARS and append a notice if truncated."""
    if len(text) <= MAX_CHARS:
        return text
    return (
        text[:MAX_CHARS]
        + f"\n\n[Document truncated at {MAX_CHARS:,} characters. "
        "Upload a shorter file or use RAG ingestion for large documents.]"
    )
